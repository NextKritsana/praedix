from __future__ import annotations

import argparse
import csv
import hashlib
import re
import sys
import zipfile
from datetime import datetime
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from pypdf import PdfReader


SOURCE_ROOT = Path(r"E:\ProJect\ACS File\advance")
OUTPUT_ROOT = Path(r"E:\ProJect\Praedix\vault\30_Knowledge_Base\ACS_Advanced\Advance_Folders")

COURSE_FOLDERS = [
    "Digital Forensic",
    "Incident Response",
    "Malware DevAnalysis",
    "Real-world System Exploitation",
    "Real-world Web Exploitation",
]

VENDOR_DIRS = {
    ".git",
    "node_modules",
    "__pycache__",
    ".pytest_cache",
    ".next",
    "dist",
    "build",
}

TEXT_EXTENSIONS = {
    ".bat",
    ".cmd",
    ".conf",
    ".config",
    ".css",
    ".csv",
    ".go",
    ".html",
    ".ini",
    ".js",
    ".json",
    ".jsx",
    ".log",
    ".md",
    ".php",
    ".ps1",
    ".py",
    ".readme",
    ".sample",
    ".sql",
    ".toml",
    ".ts",
    ".tsx",
    ".txt",
    ".xml",
    ".yaml",
    ".yml",
}

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp"}
BINARY_EXTENSIONS = {".7z", ".chm", ".dd", ".exe", ".idx", ".mkv", ".mp4", ".pack", ".rev", ".zip"}
SENSITIVE_TERMS = {
    ".env",
    "api_key",
    "apikey",
    "credential",
    "key",
    "line token",
    "password",
    "secret",
    "token",
    "vpn",
    "รหัส",
}


def configure_stdout() -> None:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")


def slug(value: str, limit: int = 100) -> str:
    value = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "-", value)
    value = re.sub(r"\s+", " ", value).strip(" .")
    return (value or "untitled")[:limit].rstrip(" .")


def file_uri(path: Path) -> str:
    return path.resolve().as_uri()


def short_hash(value: str) -> str:
    return hashlib.sha1(value.encode("utf-8", errors="ignore")).hexdigest()[:10]


def note_path(course_output: Path, source_path: Path, category: str) -> Path:
    relative = source_path.relative_to(SOURCE_ROOT)
    stem = slug(source_path.stem)
    digest = short_hash(str(relative).lower())
    return course_output / category / f"{stem}-{digest}.md"


def yaml_escape(value: str) -> str:
    return value.replace("\\", "\\\\").replace('"', '\\"')


def frontmatter(source_path: Path, kind: str, course: str) -> list[str]:
    stat = source_path.stat()
    return [
        "---",
        f'title: "{yaml_escape(source_path.stem)}"',
        f'type: "{kind}"',
        'course: "ACS Advanced"',
        f'course_folder: "{yaml_escape(course)}"',
        f'source_path: "{yaml_escape(str(source_path))}"',
        f"source_size_bytes: {stat.st_size}",
        f"source_modified: {datetime.fromtimestamp(stat.st_mtime).isoformat(timespec='seconds')}",
        f"imported_at: {datetime.now().isoformat(timespec='seconds')}",
        "tags:",
        "  - acs",
        "  - acs-advanced",
        "  - imported",
        "---",
        "",
    ]


def clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()


def read_text_file(path: Path) -> str:
    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-16", "cp874", "cp1252"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def is_sensitive(path: Path) -> bool:
    lowered = str(path).lower()
    return any(term in lowered for term in SENSITIVE_TERMS)


def is_vendor_path(path: Path) -> bool:
    return any(part in VENDOR_DIRS for part in path.parts)


def extract_pdf(source_path: Path, output_path: Path, course: str) -> tuple[int, int, str | None]:
    lines = frontmatter(source_path, "acs-advance-pdf", course)
    lines.extend([
        f"# {source_path.stem}",
        "",
        f"- Source: [{source_path.name}]({file_uri(source_path)})",
        "",
        "> Imported from PDF for Obsidian search. Verify formatting against the original file when exact layout matters.",
        "",
    ])
    try:
        reader = PdfReader(str(source_path))
        page_count = len(reader.pages)
        extracted_pages = 0
        for index, page in enumerate(reader.pages, start=1):
            try:
                text = clean_text(page.extract_text() or "")
            except Exception as exc:
                text = f"[Extraction error: {exc}]"
            if text:
                extracted_pages += 1
            lines.extend([f"## Page {index}", "", text or "_No extractable text found on this page._", ""])
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
        return page_count, extracted_pages, None
    except Exception as exc:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        lines.extend(["## Import Error", "", str(exc), ""])
        output_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
        return 0, 0, str(exc)


def extract_docx(source_path: Path, output_path: Path, course: str) -> tuple[int, str | None]:
    lines = frontmatter(source_path, "acs-advance-docx", course)
    lines.extend([
        f"# {source_path.stem}",
        "",
        f"- Source: [{source_path.name}]({file_uri(source_path)})",
        "",
    ])
    try:
        doc = Document(str(source_path))
        paragraph_count = 0
        for paragraph in doc.paragraphs:
            text = clean_text(paragraph.text)
            if text:
                paragraph_count += 1
                lines.append(text)
                lines.append("")
        for table_index, table in enumerate(doc.tables, start=1):
            lines.extend([f"## Table {table_index}", ""])
            for row in table.rows:
                cells = [clean_text(cell.text).replace("\n", "<br>") for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
            lines.append("")
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
        return paragraph_count, None
    except Exception as exc:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        lines.extend(["## Import Error", "", str(exc), ""])
        output_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
        return 0, str(exc)


def write_text_note(source_path: Path, output_path: Path, course: str) -> tuple[int, str | None]:
    lines = frontmatter(source_path, "acs-advance-text", course)
    lines.extend([
        f"# {source_path.stem}",
        "",
        f"- Source: [{source_path.name}]({file_uri(source_path)})",
        "",
    ])
    try:
        if is_sensitive(source_path):
            lines.extend([
                "## Content",
                "",
                "_Content intentionally redacted because the filename/path indicates secrets, tokens, keys, passwords, or VPN credentials._",
                "",
            ])
            line_count = 0
        else:
            text = read_text_file(source_path)
            line_count = len(text.splitlines())
            fence = source_path.suffix.lower().lstrip(".") or "text"
            lines.extend(["## Content", "", f"```{fence}", text.rstrip(), "```", ""])
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
        return line_count, None
    except Exception as exc:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        lines.extend(["## Import Error", "", str(exc), ""])
        output_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
        return 0, str(exc)


def write_asset_note(source_path: Path, output_path: Path, course: str, kind: str) -> tuple[int, str | None]:
    lines = frontmatter(source_path, kind, course)
    lines.extend([
        f"# {source_path.stem}",
        "",
        f"- Source: [{source_path.name}]({file_uri(source_path)})",
        f"- Original path: `{source_path}`",
        "",
    ])
    if source_path.suffix.lower() in IMAGE_EXTENSIONS:
        lines.extend(["## Preview", "", f"![{source_path.name}]({file_uri(source_path)})", ""])
    else:
        lines.extend(["## Binary Asset", "", "_The original binary is linked above. It was not copied into the vault._", ""])
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    return 1, None


def inventory_row(source_path: Path, course: str, status: str, note: Path | None, detail: str = "") -> dict[str, str]:
    relative = source_path.relative_to(SOURCE_ROOT / course)
    return {
        "course": course,
        "relative_path": str(relative),
        "extension": source_path.suffix.lower() or "(none)",
        "size_bytes": str(source_path.stat().st_size),
        "status": status,
        "note": str(note.relative_to(OUTPUT_ROOT)) if note else "",
        "detail": detail,
    }


def import_course(course: str) -> dict[str, int]:
    source_dir = SOURCE_ROOT / course
    course_output = OUTPUT_ROOT / slug(course)
    course_output.mkdir(parents=True, exist_ok=True)
    rows: list[dict[str, str]] = []
    stats = {
        "files": 0,
        "pdf_notes": 0,
        "docx_notes": 0,
        "text_notes": 0,
        "asset_notes": 0,
        "indexed_only": 0,
        "errors": 0,
    }

    for source_path in sorted(source_dir.rglob("*")):
        if not source_path.is_file():
            continue
        stats["files"] += 1
        suffix = source_path.suffix.lower()
        note: Path | None = None
        status = "indexed"
        detail = ""

        if is_vendor_path(source_path.relative_to(source_dir)):
            stats["indexed_only"] += 1
            rows.append(inventory_row(source_path, course, "indexed_vendor_only", None, "Skipped content import for vendor/cache/git path."))
            continue

        if suffix == ".pdf":
            note = note_path(course_output, source_path, "PDF")
            page_count, extracted_pages, error = extract_pdf(source_path, note, course)
            stats["pdf_notes"] += 1
            status = "pdf_imported"
            detail = f"{extracted_pages}/{page_count} pages with text"
            if error:
                stats["errors"] += 1
                detail = error
        elif suffix == ".docx":
            note = note_path(course_output, source_path, "DOCX")
            paragraph_count, error = extract_docx(source_path, note, course)
            stats["docx_notes"] += 1
            status = "docx_imported"
            detail = f"{paragraph_count} paragraphs"
            if error:
                stats["errors"] += 1
                detail = error
        elif suffix in TEXT_EXTENSIONS or suffix == "" or source_path.name.lower() in {"readme", "license"}:
            if source_path.stat().st_size > 1_000_000:
                note = note_path(course_output, source_path, "Assets")
                write_asset_note(source_path, note, course, "acs-advance-large-text-asset")
                stats["asset_notes"] += 1
                status = "large_text_linked"
                detail = "Large text-like file linked, not embedded."
            else:
                note = note_path(course_output, source_path, "Text")
                line_count, error = write_text_note(source_path, note, course)
                stats["text_notes"] += 1
                status = "text_imported_redacted" if is_sensitive(source_path) else "text_imported"
                detail = f"{line_count} lines"
                if error:
                    stats["errors"] += 1
                    detail = error
        elif suffix in IMAGE_EXTENSIONS or suffix in BINARY_EXTENSIONS:
            note = note_path(course_output, source_path, "Assets")
            write_asset_note(source_path, note, course, "acs-advance-asset")
            stats["asset_notes"] += 1
            status = "asset_linked"
            detail = "Linked original asset"
        else:
            stats["indexed_only"] += 1
            rows.append(inventory_row(source_path, course, "indexed_unknown_only", None, "Unknown extension; indexed only."))
            continue

        rows.append(inventory_row(source_path, course, status, note, detail))

    write_course_index(course, course_output, rows, stats)
    write_inventory_csv(course_output, rows)
    return stats


def write_inventory_csv(course_output: Path, rows: list[dict[str, str]]) -> None:
    path = course_output / "_inventory.csv"
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["course", "relative_path", "extension", "size_bytes", "status", "note", "detail"])
        writer.writeheader()
        writer.writerows(rows)


def write_course_index(course: str, course_output: Path, rows: list[dict[str, str]], stats: dict[str, int]) -> None:
    lines = [
        "---",
        f'title: "{yaml_escape(course)}"',
        "type: acs-advance-folder-index",
        'course: "ACS Advanced"',
        f'course_folder: "{yaml_escape(course)}"',
        "tags:",
        "  - acs",
        "  - acs-advanced",
        "  - index",
        "---",
        "",
        f"# {course}",
        "",
        f"- Source folder: `{SOURCE_ROOT / course}`",
        f"- Inventory CSV: [[_inventory.csv]]",
        f"- Total files seen: {stats['files']}",
        f"- PDF notes: {stats['pdf_notes']}",
        f"- DOCX notes: {stats['docx_notes']}",
        f"- Text/code notes: {stats['text_notes']}",
        f"- Asset link notes: {stats['asset_notes']}",
        f"- Indexed only: {stats['indexed_only']}",
        f"- Import errors: {stats['errors']}",
        "",
        "## Imported Notes",
        "",
        "| Source | Status | Note | Detail |",
        "| --- | --- | --- | --- |",
    ]
    for row in rows:
        if not row["note"]:
            continue
        note_stem = Path(row["note"]).stem
        source = row["relative_path"].replace("|", "\\|")
        detail = row["detail"].replace("|", "\\|")
        lines.append(f"| `{source}` | {row['status']} | [[{note_stem}]] | {detail} |")
    lines.extend([
        "",
        "## Indexed Only",
        "",
        "Files below are present in the source folder but were not embedded as content notes, usually because they are vendor/cache/git paths or unknown generated files. See `_inventory.csv` for the full list.",
        "",
    ])
    indexed = [row for row in rows if not row["note"]]
    for row in indexed[:200]:
        lines.append(f"- `{row['relative_path']}` - {row['status']}")
    if len(indexed) > 200:
        lines.append(f"- ... {len(indexed) - 200} more rows in `_inventory.csv`")
    lines.append("")
    (course_output / "_index.md").write_text("\n".join(lines), encoding="utf-8")


def write_master_index(course_stats: dict[str, dict[str, int]]) -> None:
    lines = [
        "---",
        'title: "ACS Advanced Folder Index"',
        "type: index",
        'course: "ACS Advanced"',
        "tags:",
        "  - acs",
        "  - acs-advanced",
        "  - index",
        "---",
        "",
        "# ACS Advanced Folder Index",
        "",
        "| Folder | Files seen | PDF | DOCX | Text/code | Assets | Indexed only | Errors |",
        "| --- | ---: | ---: | ---: | ---: | ---: | ---: | ---: |",
    ]
    for course, stats in course_stats.items():
        link = slug(course)
        lines.append(
            f"| [[{link}/_index|{course}]] | {stats['files']} | {stats['pdf_notes']} | {stats['docx_notes']} | "
            f"{stats['text_notes']} | {stats['asset_notes']} | {stats['indexed_only']} | {stats['errors']} |"
        )
    lines.append("")
    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)
    (OUTPUT_ROOT / "_ACS Advanced Folder Index.md").write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    configure_stdout()
    parser = argparse.ArgumentParser()
    parser.add_argument("courses", nargs="*", help="Course folder names to import. Defaults to all five folders.")
    args = parser.parse_args()

    courses = args.courses or COURSE_FOLDERS
    course_stats: dict[str, dict[str, int]] = {}
    for course in courses:
        if course not in COURSE_FOLDERS:
            raise SystemExit(f"Unknown course folder: {course}")
        print(f"Importing {course}...")
        stats = import_course(course)
        course_stats[course] = stats
        print(
            f"- files={stats['files']} pdf={stats['pdf_notes']} docx={stats['docx_notes']} "
            f"text={stats['text_notes']} assets={stats['asset_notes']} indexed_only={stats['indexed_only']} errors={stats['errors']}"
        )

    existing_stats: dict[str, dict[str, int]] = {}
    for course in COURSE_FOLDERS:
        index_path = OUTPUT_ROOT / slug(course) / "_inventory.csv"
        if index_path.exists() and course not in course_stats:
            with index_path.open("r", encoding="utf-8-sig", newline="") as handle:
                rows = list(csv.DictReader(handle))
            existing_stats[course] = {
                "files": len(rows),
                "pdf_notes": sum(1 for row in rows if row["status"].startswith("pdf_")),
                "docx_notes": sum(1 for row in rows if row["status"].startswith("docx_")),
                "text_notes": sum(1 for row in rows if row["status"].startswith("text_")),
                "asset_notes": sum(1 for row in rows if row["status"].endswith("_linked") or row["status"] == "asset_linked"),
                "indexed_only": sum(1 for row in rows if row["status"].startswith("indexed_")),
                "errors": sum(1 for row in rows if "error" in row["detail"].lower()),
            }
    write_master_index({**existing_stats, **course_stats})
    print(f"Done. Output: {OUTPUT_ROOT}")


if __name__ == "__main__":
    main()
